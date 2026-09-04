from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


NAVY = "16324F"
BLUE = "2E6F9E"
PALE_BLUE = "EAF2F7"
PALE_GRAY = "F3F5F7"
MID_GRAY = "D7DDE2"
INK = "263238"
MUTED = "68747D"
ORANGE = "D9893D"


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def set_east_asia_font(run, font_name: str):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def add_formatted_runs(paragraph, text: str, font="宋体", size=10.5):
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        bold = part.startswith("**") and part.endswith("**")
        content = part[2:-2] if bold else part
        run = paragraph.add_run(content)
        run.bold = bold
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor.from_string(INK)
        set_east_asia_font(run, "黑体" if bold else font)


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.35)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.6)
    section.right_margin = Cm(2.35)
    section.header_distance = Cm(1.1)
    section.footer_distance = Cm(1.1)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(4)
    pf.first_line_indent = Cm(0.74)
    pf.widow_control = True

    for name, size, color, before, after in [
        ("Title", 25, NAVY, 0, 12),
        ("Heading 1", 17, NAVY, 18, 10),
        ("Heading 2", 13, BLUE, 14, 6),
        ("Heading 3", 11, INK, 10, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = doc.styles["Caption"]
    caption.font.name = "宋体"
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    caption.font.size = Pt(9)
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(8)

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("GestureGraph Lab · 小组技术报告")
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    set_east_asia_font(run, "微软雅黑")
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), MID_GRAY)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)
    add_page_number(section.footer.paragraphs[0])
    section.first_page_header.paragraphs[0].text = ""
    section.first_page_footer.paragraphs[0].text = ""


def add_cover(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("GESTUREGRAPH LAB")
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(ORANGE)
    set_east_asia_font(r, "Arial")

    doc.add_paragraph("")
    doc.add_paragraph("")
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(12)
    r = title.add_run("基于时空图学习的\n14 类手势识别与渐进决策")
    r.font.size = Pt(27)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY)
    set_east_asia_font(r, "黑体")

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(28)
    r = subtitle.add_run("小组课程项目技术报告")
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor.from_string(BLUE)
    set_east_asia_font(r, "微软雅黑")

    table = doc.add_table(rows=6, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.columns[0].width = Cm(3.0)
    table.columns[1].width = Cm(11.0)
    data = [
        ("课程名称", "[课程名称]"),
        ("小组名称", "[小组名称]"),
        ("成员", "[成员 1]  [成员 2]  [成员 3]  [本人]"),
        ("任务", "22 关节点三维序列 · 正式 14 类分类"),
        ("核心路线", "ST-GCN → 双 support → AGCRN → 类别扩散"),
        ("提交日期", "2026 年 9 月 2 日（提交前可修改）"),
    ]
    for i, (k, v) in enumerate(data):
        left, right = table.rows[i].cells
        set_cell_shading(left, NAVY if i < 3 else BLUE)
        set_cell_shading(right, PALE_GRAY if i % 2 == 0 else "FFFFFF")
        for cell in (left, right):
            set_cell_margins(cell, top=110, bottom=110, start=130, end=130)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p1 = left.paragraphs[0]
        p1.paragraph_format.first_line_indent = None
        r1 = p1.add_run(k)
        r1.bold = True
        r1.font.color.rgb = RGBColor(255, 255, 255)
        r1.font.size = Pt(10.5)
        set_east_asia_font(r1, "黑体")
        p2 = right.paragraphs[0]
        p2.paragraph_format.first_line_indent = None
        r2 = p2.add_run(v)
        r2.font.size = Pt(10.5)
        set_east_asia_font(r2, "宋体")

    doc.add_paragraph("")
    note = doc.add_paragraph()
    note.paragraph_format.first_line_indent = None
    note.paragraph_format.space_before = Pt(18)
    r = note.add_run("报告原则：统一协议、逐层消融、证据分级、边界透明")
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(ORANGE)
    set_east_asia_font(r, "微软雅黑")
    doc.add_page_break()


def add_table(doc: Document, rows: list[list[str]]):
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_repeat_table_header(table.rows[0])
    for i, row in enumerate(rows):
        prevent_row_split(table.rows[i])
        for j in range(ncols):
            cell = table.cell(i, j)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_shading(cell, NAVY if i == 0 else (PALE_BLUE if i % 2 else "FFFFFF"))
            p = cell.paragraphs[0]
            p.paragraph_format.first_line_indent = None
            p.paragraph_format.line_spacing = 1.08
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
            text = row[j] if j < len(row) else ""
            text = text.replace("**", "")
            r = p.add_run(text)
            r.font.size = Pt(8.6)
            r.font.bold = i == 0
            r.font.color.rgb = RGBColor(255, 255, 255) if i == 0 else RGBColor.from_string(INK)
            set_east_asia_font(r, "黑体" if i == 0 else "宋体")
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_image(doc: Document, source_dir: Path, alt: str, relpath: str):
    path = (source_dir / relpath).resolve()
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = None
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    width = Inches(6.45)
    if path.name.startswith("equation_"):
        width = Inches(5.45)
    if "adjacency" in path.name:
        width = Inches(5.75)
    p.add_run().add_picture(str(path), width=width)
    if not alt:
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(3)
        return
    p.paragraph_format.keep_with_next = True
    caption = doc.add_paragraph(style="Caption")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.first_line_indent = None
    caption.add_run(alt)


def add_equation(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text.replace("\\operatorname", "operatorname"))
    r.font.name = "Cambria Math"
    r.font.size = Pt(11)
    r.font.italic = True


def parse_markdown(doc: Document, source_path: Path):
    lines = source_path.read_text(encoding="utf-8").splitlines()
    i = 0
    para_buf: list[str] = []
    previous_was_pagebreak = False

    def flush_paragraph():
        nonlocal para_buf
        if para_buf:
            preserve_breaks = any(s.endswith("  ") for s in para_buf)
            separator = "\n" if preserve_breaks else " "
            text = separator.join(s.strip() for s in para_buf).strip()
            p = doc.add_paragraph()
            add_formatted_runs(p, text)
            para_buf = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            flush_paragraph()
            i += 1
            continue
        if stripped == "<!-- pagebreak -->":
            flush_paragraph()
            doc.add_page_break()
            previous_was_pagebreak = True
            i += 1
            continue
        if stripped.startswith("```"):
            flush_paragraph()
            block = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                block.append(lines[i])
                i += 1
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = None
            p.paragraph_format.left_indent = Cm(0.7)
            p.paragraph_format.right_indent = Cm(0.7)
            set_cell_like = OxmlElement("w:shd")
            set_cell_like.set(qn("w:fill"), PALE_GRAY)
            p._p.get_or_add_pPr().append(set_cell_like)
            r = p.add_run("\n".join(block))
            r.font.name = "Consolas"
            r.font.size = Pt(8.5)
            i += 1
            continue
        if stripped == "$$":
            flush_paragraph()
            eq = []
            i += 1
            while i < len(lines) and lines[i].strip() != "$$":
                eq.append(lines[i].strip())
                i += 1
            add_equation(doc, " ".join(eq))
            i += 1
            continue
        image_match = re.match(r"!\[(.*?)\]\((.*?)\)", stripped)
        if image_match:
            flush_paragraph()
            add_image(doc, source_path.parent, image_match.group(1), image_match.group(2))
            i += 1
            continue
        if stripped.startswith("#"):
            flush_paragraph()
            level = len(stripped) - len(stripped.lstrip("#"))
            title = stripped[level:].strip()
            if level == 1:
                doc.add_heading(title, level=1)
            elif level == 2:
                doc.add_heading(title, level=2)
            else:
                doc.add_heading(title, level=3)
            previous_was_pagebreak = False
            i += 1
            continue
        if stripped.startswith("|") and i + 1 < len(lines) and re.match(r"^\|?\s*:?-+", lines[i + 1].strip()):
            flush_paragraph()
            table_lines = [stripped]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            rows = []
            for tl in table_lines:
                rows.append([c.strip() for c in tl.strip("|").split("|")])
            add_table(doc, rows)
            continue
        if stripped.startswith("- "):
            flush_paragraph()
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.first_line_indent = None
            add_formatted_runs(p, stripped[2:])
            i += 1
            continue
        para_buf.append(line)
        previous_was_pagebreak = False
        i += 1
    flush_paragraph()


def add_metadata(doc: Document):
    props = doc.core_properties
    props.title = "GestureGraph Lab：基于时空图学习的14类手势识别与渐进决策"
    props.subject = "小组课程项目技术报告"
    props.author = "GestureGraph Lab 小组"
    props.keywords = "ST-GCN, Graph WaveNet, AGCRN, spectral encoding, gesture recognition, class diffusion"
    props.comments = "Generated from verified local experiment reports; teammate-only results are explicitly marked."


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    args = parser.parse_args()
    doc = Document()
    configure_document(doc)
    add_metadata(doc)
    add_cover(doc)
    parse_markdown(doc, args.source.resolve())

    for section in doc.sections:
        section.start_type = WD_SECTION.NEW_PAGE
    args.output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(args.output)
    print(f"saved {args.output}")


if __name__ == "__main__":
    main()
