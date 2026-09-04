from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "handoff/GestureGraph_Lab_项目交接说明.docx"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
INK = "162033"


DOC_FONT = "冬青黑体简体中文"


def set_run_font(run, size=11, bold=False, color=INK, east_asia=DOC_FONT):
    run.font.name = east_asia
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), east_asia)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), east_asia)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = DOC_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), DOC_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), DOC_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), DOC_FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = DOC_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), DOC_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), DOC_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), DOC_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_numbering(doc, marker="bullet"):
    numbering = doc.part.numbering_part.element
    existing_abs = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    existing_num = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(existing_abs, default=0) + 1
    num_id = max(existing_num, default=0) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType"); multi.set(qn("w:val"), "singleLevel"); abstract.append(multi)
    level = OxmlElement("w:lvl"); level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start"); start.set(qn("w:val"), "1"); level.append(start)
    num_fmt = OxmlElement("w:numFmt"); num_fmt.set(qn("w:val"), "bullet" if marker == "bullet" else "decimal"); level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText"); lvl_text.set(qn("w:val"), "•" if marker == "bullet" else "%1."); level.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc"); lvl_jc.set(qn("w:val"), "left"); level.append(lvl_jc)
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs"); tab = OxmlElement("w:tab"); tab.set(qn("w:val"), "num"); tab.set(qn("w:pos"), "540"); tabs.append(tab); ppr.append(tabs)
    ind = OxmlElement("w:ind"); ind.set(qn("w:left"), "540"); ind.set(qn("w:hanging"), "270"); ppr.append(ind); level.append(ppr)
    rpr = OxmlElement("w:rPr"); fonts = OxmlElement("w:rFonts"); fonts.set(qn("w:ascii"), DOC_FONT); fonts.set(qn("w:hAnsi"), DOC_FONT); fonts.set(qn("w:eastAsia"), DOC_FONT); rpr.append(fonts); level.append(rpr)
    abstract.append(level); numbering.append(abstract)
    num = OxmlElement("w:num"); num.set(qn("w:numId"), str(num_id)); ref = OxmlElement("w:abstractNumId"); ref.set(qn("w:val"), str(abstract_id)); num.append(ref); numbering.append(num)
    return num_id


def add_list_item(doc, text, num_id):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    num_pr = paragraph._p.get_or_add_pPr().get_or_add_numPr()
    ilvl = OxmlElement("w:ilvl"); ilvl.set(qn("w:val"), "0"); num_pr.append(ilvl)
    num = OxmlElement("w:numId"); num.set(qn("w:val"), str(num_id)); num_pr.append(num)
    set_run_font(paragraph.add_run(text))
    return paragraph


def set_cell_shading(cell, fill):
    shd = cell._tc.get_or_add_tcPr().find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd"); cell._tc.get_or_add_tcPr().append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW"); tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa"); tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd"); tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa"); tbl_ind.set(qn("w:w"), "120")
    grid = table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol"); col.set(qn("w:w"), str(width)); grid.append(col)
    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar"); tbl_pr.append(margins)
    for side, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
        item = margins.find(qn(f"w:{side}"))
        if item is None:
            item = OxmlElement(f"w:{side}"); margins.append(item)
        item.set(qn("w:w"), str(value)); item.set(qn("w:type"), "dxa")
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            tc_w = cell._tc.get_or_add_tcPr().get_or_add_tcW(); tc_w.set(qn("w:type"), "dxa"); tc_w.set(qn("w:w"), str(widths[index]))
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, text in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, LIGHT_BLUE)
        paragraph = cell.paragraphs[0]; paragraph.paragraph_format.space_after = Pt(0)
        set_run_font(paragraph.add_run(text), size=10, bold=True, color=DARK_BLUE)
    for row in rows:
        cells = table.add_row().cells
        for index, text in enumerate(row):
            paragraph = cells[index].paragraphs[0]; paragraph.paragraph_format.space_after = Pt(0); paragraph.paragraph_format.line_spacing = 1.15
            set_run_font(paragraph.add_run(str(text)), size=9.5)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_callout(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0); set_cell_shading(cell, LIGHT_GRAY)
    p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(2)
    set_run_font(p.add_run(title + "  "), size=10.5, bold=True, color=DARK_BLUE)
    set_run_font(p.add_run(text), size=10.5)
    set_table_geometry(table, [9360])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = " PAGE "
    separate = OxmlElement("w:fldChar"); separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t"); text.text = "1"
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, text, end): run._r.append(node)
    set_run_font(run, size=9, color=MUTED)


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5); section.page_height = Inches(11)
    section.top_margin = section.right_margin = section.bottom_margin = section.left_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(.492)
    configure_styles(doc)
    bullet_id = add_numbering(doc, "bullet")

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(header.add_run("GestureGraph Lab  |  项目交接"), size=9, bold=True, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(footer.add_run("2026-08-25  ·  Page "), size=9, color=MUTED); add_page_field(footer)

    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(18); p.paragraph_format.space_after = Pt(4)
    set_run_font(p.add_run("项目交接说明"), size=24, bold=True, color=INK)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(16)
    set_run_font(p.add_run("GestureGraph Lab · SHREC'17 时空图手势识别"), size=13, color=MUTED)
    for label, value in (("交接日期", "2026-08-25"), ("当前状态", "课程主体完成；Mac 与 Windows 核心链路及摄像头演示均已实测"), ("正式模型", "ST-GCN · 14 类 · 官方测试集 70.95%")):
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
        set_run_font(p.add_run(f"{label}："), bold=True, color=DARK_BLUE); set_run_font(p.add_run(value))
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    add_callout(doc, "核心结论", "课程要求的模型、基线与消融实验已经完成。摄像头页面是额外原型；其现场准确率尚未建立，不能直接引用 70.95%。")

    doc.add_heading("1. 项目是什么", level=1)
    p = doc.add_paragraph("项目把一段手势表示为连续的 22 节点手部图，用 ST-GCN 同时学习骨架空间连接和时间运动，并与 MLP、去图结构和单帧模型进行对照。")
    p.paragraph_format.keep_with_next = True
    add_callout(doc, "数据流", "摄像头/SHREC 骨架 → 22 点标准化 → 24-64 帧动作切分 → ST-GCN → 14 类概率 → 拒识与投票 → 中文结果。")

    doc.add_heading("2. 当前完成状态", level=1)
    for item in (
        "官方 SHREC'17 14 类训练、验证和官方测试已完成。",
        "完整 ST-GCN、MLP、无图邻接 ST-GCN、单帧 ST-GCN 四组对照已完成。",
        "关节组遮挡、混淆矩阵、正确/错误动画与 Markdown 报告已完成。",
        "MediaPipe 摄像头、22 点归一化、Model Space、多视角和本地推理已完成。",
        "动作自动切分、低置信度拒识和连续预测投票已完成。",
        "页面采集标签已统一为官方 14 类，并提供中英文说明和参考动画。",
        "Windows 真机已通过依赖安装、9/9 环境检查、8/8 前端测试、5/5 Python 测试、模型/API、摄像头骨架与 Model Space 验证。",
    ): add_list_item(doc, item, bullet_id)

    doc.add_page_break()
    doc.add_heading("3. 核心实验结果", level=1)
    add_table(doc, ["实验", "官方测试集", "相对完整模型"], [
        ("完整 ST-GCN", "70.95%", "基线"),
        ("展平 MLP", "63.81%", "-7.14 pp"),
        ("去掉图结构", "68.93%", "-2.02 pp"),
        ("只看单帧", "32.02%", "-38.93 pp"),
    ], [4700, 2100, 2560])
    p = doc.add_paragraph("实验解释：时间动态是主要信号，显式骨架拓扑提供额外增益，完整 ST-GCN 优于不使用图结构的 MLP。")

    doc.add_heading("4. 快速启动", level=1)
    doc.add_heading("macOS", level=2)
    mac_steps = add_numbering(doc, "decimal")
    for item in ("准备 Python 3.10-3.12 和 Node.js 20+。", "运行 handoff/start_mac.command，打开 http://localhost:8080/ 并允许摄像头。"):
        add_list_item(doc, item, mac_steps)
    doc.add_heading("Windows 10/11", level=2)
    windows_steps = add_numbering(doc, "decimal")
    for item in ("准备 64 位 Python 3.10-3.12 和 Node.js 20+。", "运行 handoff\\start_windows.bat，使用 Edge/Chrome 打开 http://localhost:8080/。"):
        add_list_item(doc, item, windows_steps)
    add_callout(doc, "检查命令", "npm run check · npm test · npm run test:python")

    doc.add_heading("5. Mac / Windows 兼容性", level=1)
    add_table(doc, ["模块", "macOS", "Windows"], [
        ("Python 推理/训练", "Apple Silicon MPS/CPU 已实测", "3.10.11 / PyTorch CPU 已实测"),
        ("Node/npm", "已实测", "Node 24.15；9/9、8/8、5/5 通过"),
        ("摄像头页面", "Safari 已实测", "Edge/Chrome 摄像头、骨架与 Model Space 已实测"),
        ("启动流程", "start_mac.command 已实测", "命令行已实测；bat 双击待确认"),
        ("正式 14 类模型", "已加载并推理", "HandSTGCN / 14 类预测已实测"),
    ], [2550, 3300, 3510])

    doc.add_heading("6. 已知限制与风险", level=1)
    for item in (
        "70.95% 是 SHREC'17 官方测试集准确率，不等于普通摄像头准确率。",
        "官方深度传感器与 MediaPipe 摄像头存在数据域差异。",
        "摄像头 14 类尚未建立独立准确率和混淆矩阵。",
        "左右挥动与顺/逆时针需要专项检查镜像语义。",
        "动作切分阈值仅根据少量本机数据设定，尚未针对多人校准。",
        "保存样本只写入 JSON，不会自动训练模型。",
        "Windows 验证副本已修复 AppleDouble 文件导致参考 GIF 误计数的问题；该修改仍需同步回主交接包。",
        "原工作区的旧 open_palm 样本属于早期 6 类体系；轻量交接包不复制任何个人录制样本。",
        "完整 SHREC'17 数据约 6 GB，不进入轻量交接包；演示不需要它。",
    ): add_list_item(doc, item, bullet_id)

    doc.add_heading("7. 关键目录", level=1)
    add_table(doc, ["路径", "用途"], [
        ("index.html / styles.css / js", "摄像头界面、切分、投票和 Model Space"),
        ("gesturegraph/model.py", "ST-GCN 与 MLP"),
        ("gesturegraph/train.py", "训练、验证和官方测试"),
        ("gesturegraph/serve.py", "本地页面、样本保存和推理 API"),
        ("runs/.../stgcn_full/best.pt", "当前 14 类正式模型"),
        ("runs/.../REPORT.md", "实验报告与结论"),
        ("assets/models", "MediaPipe 浏览器手部检测模型"),
        ("assets/references", "14 类官方骨架参考动画"),
        ("data/recordings", "新采集的官方 14 类 JSON"),
    ], [3550, 5810])

    doc.add_heading("8. 常见问题回答", level=1)
    qa = (
        ("为什么是 64 帧？", "统一不同长度动作；现场自动捕获 24-64 帧，服务端再重采样。"),
        ("保存样本会自动训练吗？", "不会，只保存带标签的 22 点坐标 JSON。"),
        ("为什么有未识别？", "低置信度或多次预测不一致时拒识，比强行猜测可靠。"),
        ("项目完成了吗？", "课程官方实验完成；摄像头产品化仍需测试集与微调。"),
        ("Windows 完全支持吗？", "Windows 10/11 CPU 环境、服务、模型、API 与摄像头演示均已通过真机验证。"),
    )
    for question, answer in qa:
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(5)
        set_run_font(p.add_run(question + "  "), bold=True, color=DARK_BLUE); set_run_font(p.add_run(answer))

    doc.add_heading("9. 交接验收清单", level=1)
    for item in (
        "Windows 环境检查 9/9、前端测试 8/8、Python 测试 5/5 已通过。",
        "Windows 服务、/api/health、14 类预测 API 和隔离目录样本保存已通过。",
        "Windows 浏览器摄像头、22 点骨架和 Model Space 已通过人工验收。",
        "动作状态能经历采集中、分析中、识别或未识别。",
        "下拉框包含 14 类并能切换参考动画。",
        "组员能找到模型、报告、结果和样本目录。",
        "组员理解官方数据准确率与摄像头准确率不同。",
    ): add_list_item(doc, "□ " + item, bullet_id)

    doc.add_heading("10. 建议下一阶段", level=1)
    next_steps = add_numbering(doc, "decimal")
    for item in (
        "把 Windows 验证副本中的 AppleDouble GIF 过滤修复同步回主交接包。",
        "建立摄像头 14 类测试模式，每类 10-20 段并生成混淆矩阵。",
        "检查左右挥动与旋转方向的镜像定义。",
        "校准动作切分和每类拒识阈值。",
        "采集多人 MediaPipe 数据微调模型。",
        "尝试坐标 + 速度 + 骨骼向量并与 70.95%基线比较。",
    ): add_list_item(doc, item, next_steps)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
